"""DOCX-отчёты для Заказчика (v0.11) — деловой альбомный формат.

Два типа:
  • `build_variant_report(name, tep, path)` — подробный отчёт по ОДНОМУ варианту
    (база / максимум площади / максимум прибыли и т.п.): титул, общие ТЭП,
    баланс территории (таблица + диаграмма), соцобъекты, парковки, озеленение,
    кластеры этажности, экономика (себестоимость/выручка/метрики + диаграмма).
  • `build_comparison_report(scenarios, path)` — сравнение нескольких вариантов:
    сводная таблица, диаграммы сравнения, рекомендация лучшего.

`build_report` — алиас на `build_comparison_report` (обратная совместимость).

Стиль: альбомная A4, фирменный синий, заливка шапок таблиц, колонтитулы с
номерами страниц, заголовки секций с подчёркиванием. python-docx + matplotlib.
"""

from __future__ import annotations

import datetime as _dt
from io import BytesIO

from urban_model.models.result import TEPResult

Scenario = tuple[str, TEPResult]

# --- Фирменная палитра ---
_PRIMARY = "1F4E79"      # глубокий синий — заголовки/акцент
_HEADER_FILL = "1F4E79"  # заливка шапок таблиц (текст белый)
_ZEBRA_FILL = "EAF1F8"   # светло-голубой — чётные строки
_TOTAL_FILL = "D6E4F0"   # итоговые строки
_GREY_TEXT = "595959"


# ---------------------------------------------------------------------------
# Форматирование чисел
# ---------------------------------------------------------------------------

def _fmt(v: float | int | None, nd: int = 0, dash: str = "—") -> str:
    if v is None:
        return dash
    if nd == 0:
        return f"{v:,.0f}".replace(",", " ")
    return f"{v:,.{nd}f}".replace(",", " ")


def _profit(tep: TEPResult) -> float | None:
    return float(tep.economy.profit) if tep.economy is not None else None


# ---------------------------------------------------------------------------
# Низкоуровневые helpers python-docx (XML)
# ---------------------------------------------------------------------------

def _shade_cell(cell, fill_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_cell_margins(cell, top=60, bottom=60, left=120, right=120) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("bottom", bottom), ("start", left),
                      ("end", right), ("left", left), ("right", right)):
        e = OxmlElement(f"w:{name}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def _heading_border(paragraph, color_hex: str = _PRIMARY) -> None:
    """Тонкая линия под заголовком секции."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _add_footer_page_numbers(section, project_label: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{project_label}    ·    Стр. ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    # поле PAGE
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    r2 = p.add_run(); r2.font.size = Pt(8); r2.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    r2._r.append(fld_begin); r2._r.append(instr); r2._r.append(fld_end)


def _new_landscape_doc() -> "Document":  # noqa: F821
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.shared import Cm, Pt

    doc = Document()
    # Базовый шрифт
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    # A4 альбомная: 29.7 × 21.0 см
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.3)
    return doc


def _content_width_cm() -> float:
    return 29.7 - 1.5 - 1.5  # 26.7 см


def _section_heading(doc, text: str) -> None:
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    _heading_border(p)


def _styled_table(doc, headers: list[str], rows: list[list[str]],
                  col_widths_cm: list[float] | None = None,
                  total_row_idx: int | None = None,
                  font_size: int = 9) -> None:
    """Деловая таблица: синяя шапка (белый текст), зебра, отступы.

    `total_row_idx` — индекс строки данных, выделяемой как «итог» (заливка).
    """
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Cm, Pt, RGBColor

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    # шапка
    hdr = table.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = ""
        para = hdr[j].paragraphs[0]
        r = para.add_run(h)
        r.bold = True
        r.font.size = Pt(font_size)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(hdr[j], _HEADER_FILL)
        _set_cell_margins(hdr[j])
    # строки
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        is_total = (total_row_idx is not None and i == total_row_idx)
        for j, val in enumerate(row):
            cells[j].text = ""
            para = cells[j].paragraphs[0]
            r = para.add_run(str(val))
            r.font.size = Pt(font_size)
            if is_total:
                r.bold = True
            _set_cell_margins(cells[j])
            if is_total:
                _shade_cell(cells[j], _TOTAL_FILL)
            elif i % 2 == 1:
                _shade_cell(cells[j], _ZEBRA_FILL)
    # ширины колонок
    if col_widths_cm:
        for j, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[j].width = Cm(w)
    return table


def _kpi_grid(doc, items: list[tuple[str, str]], cols: int = 4) -> None:
    """Сетка «крупное значение + подпись» — общие ТЭП на видном месте."""
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Cm, Pt, RGBColor

    n = len(items)
    rows_n = (n + cols - 1) // cols
    table = doc.add_table(rows=rows_n, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    w = _content_width_cm() / cols
    for idx, (label, value) in enumerate(items):
        r, c = divmod(idx, cols)
        cell = table.rows[r].cells[c]
        cell.text = ""
        _shade_cell(cell, _ZEBRA_FILL)
        _set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
        cell.width = Cm(w)
        pv = cell.paragraphs[0]
        rv = pv.add_run(value)
        rv.bold = True
        rv.font.size = Pt(15)
        rv.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        pl = cell.add_paragraph()
        rl = pl.add_run(label)
        rl.font.size = Pt(8)
        rl.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    # пустые ячейки последнего ряда
    for idx in range(n, rows_n * cols):
        r, c = divmod(idx, cols)
        table.rows[r].cells[c].width = Cm(w)


def _title_block(doc, title: str, subtitle: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    sub = doc.add_paragraph()
    rs = sub.add_run(subtitle)
    rs.font.size = Pt(10)
    rs.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    meta = doc.add_paragraph()
    rm = meta.add_run(
        f"Дата формирования: {_dt.date.today().strftime('%d.%m.%Y')}   ·   "
        f"Профиль нормативов: СПб"
    )
    rm.font.size = Pt(9)
    rm.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    _heading_border(meta)


# ---------------------------------------------------------------------------
# Диаграммы (matplotlib → PNG)
# ---------------------------------------------------------------------------

def _mpl():
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    plt.rcParams.update({
        "font.size": 9, "axes.titlesize": 11, "axes.edgecolor": "#888888",
        "axes.grid": True, "grid.color": "#E0E0E0", "grid.linewidth": 0.6,
    })
    return plt


_PRETTY_BALANCE = {
    "housing_lot": "ЗУ жилой застройки", "kindergarten_plot": "ДОО",
    "school_plot": "СОШ", "sport_facilities": "Спорт",
    "social_parking_plot": "Парк. соцобъектов", "znop": "ЗНОП",
    "intra_quarter_driveways": "Проезды", "parking_multilevel": "Многоур. паркинги",
    "custom_objects": "Доп. объекты", "engineering_plot": "Инж. инфраструктура",
}


def _chart_balance(tep: TEPResult) -> BytesIO | None:
    try:
        plt = _mpl()
    except Exception:  # noqa: BLE001
        return None
    b = tep.balance
    site = b.site_area
    if site <= 0:
        return None
    items = [(self_pretty, v) for k, v in sorted(b.components.items(), key=lambda kv: kv[1])
             if v > 0 for self_pretty in [_PRETTY_BALANCE.get(k, k)]]
    surplus = max(0.0, b.surplus)
    if surplus > 0:
        items.insert(0, ("Резерв (озеленение)", surplus))
    if not items:
        return None
    labels = [x[0] for x in items]
    vals = [x[1] for x in items]
    pcts = [v / site * 100 for v in vals]
    # v0.10.10: цветные столбцы (палитра) вместо однотонных — нагляднее.
    palette = [
        "#1F4E79", "#2E75B6", "#5B9BD5", "#9DC3E6", "#70AD47", "#A9D18E",
        "#F5A623", "#ED7D31", "#C0C0C0", "#BFBFBF",
    ]
    colors = [palette[i % len(palette)] for i in range(len(items))]
    # Резерв — мягкий зелёный (последний/первый элемент = озеленение)
    for i, lbl in enumerate(labels):
        if "Резерв" in lbl:
            colors[i] = "#9CCC65"
    fig, ax = plt.subplots(figsize=(7.2, 0.5 * len(items) + 1.0))
    bars = ax.barh(labels, vals, color=colors, edgecolor="white", height=0.7)
    xmax = max(vals) if vals else 1.0
    for bar, pct, v in zip(bars, pcts, vals):
        ax.text(bar.get_width() + xmax * 0.01, bar.get_y() + bar.get_height() / 2,
                f"{v:,.0f} м² · {pct:.0f}%".replace(",", " "),
                va="center", fontsize=8.5, color="#333333")
    ax.set_xlim(0, xmax * 1.28)  # место под подписи
    ax.set_xlabel("Площадь, м²")
    ax.set_title("Баланс территории квартала", fontweight="bold", color="#1F4E79")
    ax.grid(axis="y", visible=False)
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return _save(fig, plt)


def _chart_economy(tep: TEPResult) -> BytesIO | None:
    if tep.economy is None:
        return None
    try:
        plt = _mpl()
    except Exception:  # noqa: BLE001
        return None
    e = tep.economy
    # v0.10.10: простая и понятная картинка — 3 столбца:
    # Себестоимость / Выручка / Прибыль (последний может быть отрицательным).
    cost_t = e.cost.total
    rev_t = e.revenue.total
    profit = e.profit
    labels = ["Себестоимость", "Выручка", "Прибыль"]
    vals = [cost_t, rev_t, profit]
    colors = ["#C0504D", "#4CAF50", "#1F4E79" if profit >= 0 else "#C0504D"]
    fig, ax = plt.subplots(figsize=(6.6, 4.0))
    bars = ax.bar(labels, vals, color=colors, width=0.55, edgecolor="white")
    ax.axhline(0, color="#888888", linewidth=0.8)
    for bar, v in zip(bars, vals):
        ax.text(bar.get_x() + bar.get_width() / 2,
                bar.get_height() + (max(abs(x) for x in vals) * 0.02
                                    if v >= 0 else -max(abs(x) for x in vals) * 0.05),
                f"{v:+,.0f}".replace(",", " ") if v is profit else f"{v:,.0f}".replace(",", " "),
                ha="center", va="bottom" if v >= 0 else "top",
                fontsize=11, fontweight="bold", color="#222222")
    ax.set_ylabel("Баллы выгодности")
    ax.set_title("Экономика: себестоимость → выручка → прибыль",
                 fontweight="bold", color="#1F4E79")
    ax.grid(axis="x", visible=False)
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    fig.tight_layout()
    return _save(fig, plt)


def _chart_comparison(scenarios: list[Scenario]) -> BytesIO | None:
    try:
        plt = _mpl()
    except Exception:  # noqa: BLE001
        return None
    names = [n for n, _ in scenarios]
    apts = [(t.apartments_area.value or 0.0) for _, t in scenarios]
    profits = [_profit(t) for _, t in scenarios]
    has_profit = any(p is not None for p in profits)
    n = len(scenarios)
    fig, ax1 = plt.subplots(figsize=(max(7, n * 1.6), 4.2))
    x = range(n)
    width = 0.38 if has_profit else 0.6
    ax1.bar([i - (width / 2 if has_profit else 0) for i in x], apts,
            width=width, color="#1F4E79", label="Площадь квартир, м²")
    ax1.set_ylabel("Площадь квартир, м²", color="#1F4E79")
    ax1.tick_params(axis="y", labelcolor="#1F4E79")
    ax1.set_xticks(list(x))
    ax1.set_xticklabels(names, rotation=15, ha="right", fontsize=8)
    if has_profit:
        ax2 = ax1.twinx()
        ax2.bar([i + width / 2 for i in x], [p or 0.0 for p in profits],
                width=width, color="#2E7D32", label="Выгодность, баллы")
        ax2.set_ylabel("Выгодность, баллы", color="#2E7D32")
        ax2.tick_params(axis="y", labelcolor="#2E7D32")
        ax2.axhline(0, color="#999999", linewidth=0.8)
        ax2.grid(False)
    ax1.set_title("Сравнение вариантов")
    fig.tight_layout()
    return _save(fig, plt)


def _save(fig, plt) -> BytesIO:
    buf = BytesIO()
    fig.savefig(buf, format="png", dpi=145, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _add_image(doc, buf: BytesIO | None, width_cm: float) -> None:
    if buf is None:
        return
    from docx.shared import Cm
    doc.add_picture(buf, width=Cm(width_cm))


# ---------------------------------------------------------------------------
# Сбор блоков ТЭП (общие для отчёта по варианту)
# ---------------------------------------------------------------------------

def _general_kpis(tep: TEPResult) -> list[tuple[str, str]]:
    items = [
        ("КИТ (ПЗЗ)", _fmt(tep.kit.value, 3)),
        ("Площадь квартир, м²", _fmt(tep.apartments_area.value)),
        ("Население, чел.", _fmt(tep.population.value)),
        ("Плотность, чел/га", _fmt(tep.density_chel_per_ga.value, 1)),
        ("Общая площадь зданий, м²", _fmt(tep.gfa.value)),
        ("Парковки, м/м", _fmt(tep.parking_required_places.value)),
        ("ЗНОП, м²/чел", _fmt(tep.znop_per_person.value, 1)),
        ("Резерв баланса, м²", _fmt(tep.balance.surplus)),
    ]
    if tep.economy is not None:
        e = tep.economy
        items += [
            ("Выгодность, баллы", _fmt(e.profit)),
            ("Маржа, %", f"{e.margin * 100:.1f}"),
            ("ROI, %", f"{e.roi * 100:.1f}"),
            ("Выход жилья, %", f"{e.sellable_ratio * 100:.0f}"),
        ]
    return items


def _balance_rows(tep: TEPResult) -> list[list[str]]:
    b = tep.balance
    site = b.site_area or 1.0
    rows = []
    for k, v in sorted(b.components.items(), key=lambda kv: -kv[1]):
        if v <= 0:
            continue
        rows.append([_PRETTY_BALANCE.get(k, k), _fmt(v), f"{v / site * 100:.1f}%"])
    rows.append(["Итого занято", _fmt(b.required_total), f"{b.required_total / site * 100:.1f}%"])
    rows.append(["Резерв (озеленение/двор)", _fmt(b.surplus), f"{b.surplus / site * 100:.1f}%"])
    return rows


# ---------------------------------------------------------------------------
# Отчёт по ОДНОМУ варианту
# ---------------------------------------------------------------------------

def build_variant_report(name: str, tep: TEPResult, path: str) -> str:
    """Подробный деловой отчёт по одному варианту застройки."""
    from docx.shared import Cm

    site_area = tep.balance.site_area
    doc = _new_landscape_doc()
    _add_footer_page_numbers(doc.sections[0], f"Отчёт: {name}")

    _title_block(
        doc,
        "Технико-экономические показатели застройки",
        f"Вариант: «{name}»   ·   Площадь квартала: "
        f"{_fmt(site_area)} м² ({site_area / 10000:.2f} га)",
    )

    # 1. Общие ТЭП
    _section_heading(doc, "Основные показатели")
    _kpi_grid(doc, _general_kpis(tep), cols=4)

    # 2. Баланс территории — таблица + диаграмма (две колонки)
    _section_heading(doc, "Баланс территории")
    cw = _content_width_cm()
    tbl = doc.add_table(rows=1, cols=2)
    left, right = tbl.rows[0].cells
    left.width = Cm(cw * 0.5)
    right.width = Cm(cw * 0.5)
    # таблица слева
    _table_into_cell(left, ["Компонент", "Площадь, м²", "Доля"],
                     _balance_rows(tep), total_row_idx=None)
    # диаграмма справа
    _img_into_cell(right, _chart_balance(tep), cw * 0.48)

    # 3. Жильё / соцобъекты / парковки / озеленение
    _section_heading(doc, "Жильё и социальная инфраструктура")
    _styled_table(
        doc,
        ["Показатель", "Значение", "Ед."],
        [
            ["Площадь квартир", _fmt(tep.apartments_area.value), "м²"],
            ["Общая площадь зданий (GFA)", _fmt(tep.gfa.value), "м²"],
            ["ВПП (встроенно-пристроенные)", _fmt(tep.built_in_area.value), "м²"],
            ["ДОО — требуется / принято",
             f"{_fmt(tep.kindergarten_places_required.value)} / {_fmt(tep.kindergarten_places_accepted.value)}", "мест"],
            ["СОШ — требуется / принято",
             f"{_fmt(tep.school_places_required.value)} / {_fmt(tep.school_places_accepted.value)}", "мест"],
            ["ЗНОП", _fmt(tep.znop_area.value), "м²"],
            ["Озеленение жилого ЗУ", _fmt(tep.greening_housing_area.value), "м²"],
        ],
        col_widths_cm=[cw * 0.5, cw * 0.3, cw * 0.2],
    )

    _section_heading(doc, "Парковки")
    _styled_table(
        doc,
        ["Тип", "Машино-мест", "Площадь, м²"],
        [
            ["Всего требуется", _fmt(tep.parking_required_places.value), "—"],
            ["Открытые наземные", _fmt(tep.parking_open_places.value), _fmt(tep.parking_open_area.value)],
            ["Многоуровневые", _fmt(tep.parking_multilevel_places.value), _fmt(tep.parking_multilevel_area.value)],
            ["Подземные", _fmt(tep.parking_underground_places.value), "не на поверхности"],
            ["Парковки соцобъектов", _fmt(tep.social_parking_total.value), _fmt(tep.social_parking_area.value)],
        ],
        col_widths_cm=[cw * 0.4, cw * 0.3, cw * 0.3],
    )

    # 4. Кластеры этажности (если заданы)
    if tep.floor_clusters_detail:
        _section_heading(doc, "Кластеры этажности (по зонам)")
        rows = [
            [d["label"], _fmt(d["area_m2"]), str(d["floors"]),
             _fmt(d["kit"], 3), _fmt(d["apartments_area"]), _fmt(d["footprint"])]
            for d in tep.floor_clusters_detail
        ]
        _styled_table(
            doc,
            ["Зона", "Площадь, м²", "Этажей", "КИТ зоны", "Кварт., м²", "Пятно, м²"],
            rows,
            col_widths_cm=[cw * 0.22, cw * 0.18, cw * 0.13, cw * 0.15, cw * 0.16, cw * 0.16],
        )

    # 5. Экономика
    if tep.economy is not None:
        e = tep.economy
        _section_heading(doc, "Экономическая оценка (баллы выгодности)")
        eco_tbl = doc.add_table(rows=1, cols=2)
        ec_left, ec_right = eco_tbl.rows[0].cells
        ec_left.width = Cm(cw * 0.5)
        ec_right.width = Cm(cw * 0.5)
        cb, rb = e.cost, e.revenue
        _table_into_cell(
            ec_left, ["Статья", "Себестоимость", "Выручка"],
            [
                ["Жильё", _fmt(cb.residential), _fmt(rb.residential)],
                ["ВПП / коммерция", _fmt(cb.vpp), _fmt(rb.vpp_commercial + rb.custom_commercial)],
                ["ДОО + СОШ", _fmt(cb.kindergarten + cb.school), _fmt(rb.social_compensation)],
                ["Парковки", _fmt(cb.parking_open + cb.parking_multilevel + cb.parking_underground),
                 _fmt(rb.parking_open + rb.parking_multilevel + rb.parking_underground)],
                ["Накладные", _fmt(cb.networks + cb.landscaping + cb.design + cb.contingency), "—"],
                ["ИТОГО", _fmt(cb.total), _fmt(rb.total)],
            ],
            total_row_idx=5,
        )
        _img_into_cell(ec_right, _chart_economy(tep), cw * 0.48)
        _kpi_grid(doc, [
            ("Выгодность, баллы", _fmt(e.profit)),
            ("Прибыль без соц.", _fmt(e.profit_before_social)),
            ("Соц. нагрузка", _fmt(-e.net_social_burden)),
            ("Маржа / ROI", f"{e.margin*100:.1f}% / {e.roi*100:.1f}%"),
        ], cols=4)

    _footnote(doc)
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# Вспомогательные: таблица/картинка ВНУТРЬ ячейки (для 2-колоночной вёрстки)
# ---------------------------------------------------------------------------

def _table_into_cell(cell, headers, rows, total_row_idx=None) -> None:
    from docx.shared import Pt, RGBColor
    # очищаем дефолтный пустой параграф
    cell.text = ""
    inner = cell.add_table(rows=1, cols=len(headers))
    inner.style = "Table Grid"
    hdr = inner.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = ""
        r = hdr[j].paragraphs[0].add_run(h)
        r.bold = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(hdr[j], _HEADER_FILL); _set_cell_margins(hdr[j])
    for i, row in enumerate(rows):
        cells = inner.add_row().cells
        is_total = (total_row_idx is not None and i == total_row_idx)
        for j, val in enumerate(row):
            cells[j].text = ""
            r = cells[j].paragraphs[0].add_run(str(val))
            r.font.size = Pt(8)
            if is_total:
                r.bold = True
            _set_cell_margins(cells[j])
            if is_total:
                _shade_cell(cells[j], _TOTAL_FILL)
            elif i % 2 == 1:
                _shade_cell(cells[j], _ZEBRA_FILL)


def _img_into_cell(cell, buf, width_cm) -> None:
    from docx.shared import Cm
    if buf is None:
        return
    cell.text = ""
    run = cell.paragraphs[0].add_run()
    run.add_picture(buf, width=Cm(width_cm))


def _footnote(doc) -> None:
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    r = p.add_run(
        "Примечание: «баллы выгодности» — безразмерный индикатор для сравнения "
        "вариантов (1.0 ≈ м² жилья 9-эт. монолита со standard-отделкой), не "
        "денежная оценка. Расчёт по нормативам ПЗЗ/СП СПб."
    )
    r.italic = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)


# ---------------------------------------------------------------------------
# Отчёт-СРАВНЕНИЕ
# ---------------------------------------------------------------------------

def _pick_best(scenarios: list[Scenario]) -> tuple[int, str]:
    feasible = [(i, n, t) for i, (n, t) in enumerate(scenarios) if t.balance.is_feasible]
    pool = feasible if feasible else [(i, n, t) for i, (n, t) in enumerate(scenarios)]
    with_econ = [(i, n, t) for i, n, t in pool if t.economy is not None]
    if with_econ:
        i, name, _ = max(with_econ, key=lambda it: it[2].economy.profit)
        basis = "максимальная выгодность среди вариантов с балансом"
    else:
        i, name, _ = max(pool, key=lambda it: (it[2].apartments_area.value or 0.0))
        basis = "максимальная площадь квартир среди вариантов с балансом"
    if not feasible:
        basis += " (внимание: ни один вариант не проходит баланс полностью)"
    return i, f"«{name}» — {basis}."


def build_comparison_report(scenarios: list[Scenario], path: str) -> str:
    """Деловой отчёт-сравнение нескольких вариантов."""
    from docx.shared import Cm

    if not scenarios:
        raise ValueError("Нет сценариев для отчёта.")

    doc = _new_landscape_doc()
    _add_footer_page_numbers(doc.sections[0], "Сравнение вариантов застройки")
    _title_block(
        doc,
        "Сравнение вариантов застройки территории",
        f"Вариантов в сравнении: {len(scenarios)}",
    )

    # Сводная таблица
    _section_heading(doc, "Сводная таблица ТЭП")
    headers = ["Вариант", "КИТ", "Площ. кв., м²", "Нас.", "ДОО", "СОШ",
               "Парк., м/м", "ЗНОП", "Резерв, м²", "Выгодн.", "Баланс"]
    rows = []
    for nm, t in scenarios:
        rows.append([
            nm, _fmt(t.kit.value, 3), _fmt(t.apartments_area.value),
            _fmt(t.population.value), _fmt(t.kindergarten_places_accepted.value),
            _fmt(t.school_places_accepted.value), _fmt(t.parking_required_places.value),
            _fmt(t.znop_per_person.value, 1), _fmt(t.balance.surplus),
            (_fmt(_profit(t)) if t.economy is not None else "—"),
            ("✓" if t.balance.is_feasible else "дефицит"),
        ])
    cw = _content_width_cm()
    _styled_table(doc, headers, rows, font_size=8)

    # Диаграмма сравнения
    _section_heading(doc, "Диаграмма сравнения")
    _add_image(doc, _chart_comparison(scenarios), min(cw, 24.0))

    # Рекомендация
    _section_heading(doc, "Рекомендация")
    best_idx, basis = _pick_best(scenarios)
    best_name, best_tep = scenarios[best_idx]
    pr = doc.add_paragraph()
    from docx.shared import Pt
    rr = pr.add_run(f"Рекомендуемый вариант: {basis}")
    rr.bold = True; rr.font.size = Pt(11)
    _kpi_grid(doc, _general_kpis(best_tep), cols=4)

    _footnote(doc)
    doc.save(path)
    return path


# Обратная совместимость: build_report == сравнение.
def build_report(scenarios: list[Scenario], path: str) -> str:
    return build_comparison_report(scenarios, path)
