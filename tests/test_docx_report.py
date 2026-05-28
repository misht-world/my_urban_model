"""Тесты v0.10 — DOCX-отчёт по сценариям."""

from __future__ import annotations

import os
import tempfile

import pytest

from urban_model import solve_max_kit
from urban_model.export.docx_report import _pick_best, build_report
from urban_model.models import CalculationOptions, Site
from urban_model.normatives import load_normatives


@pytest.fixture(scope="module")
def spb():
    return load_normatives("spb")


@pytest.fixture(scope="module")
def scenarios(spb):
    s = Site(area_m2=50_000)
    return [
        ("12 эт комфорт", solve_max_kit(s, CalculationOptions(floors=12, residential_class="comfort"), spb)),
        ("18 эт бизнес", solve_max_kit(s, CalculationOptions(floors=18, residential_class="business"), spb)),
    ]


def test_build_report_creates_openable_docx(scenarios):
    path = tempfile.mktemp(suffix=".docx")
    try:
        build_report(scenarios, path)
        assert os.path.getsize(path) > 5000  # непустой, с диаграммой
        from docx import Document
        doc = Document(path)
        # v0.11: сводная таблица + KPI-сетка рекомендации.
        assert len(doc.tables) >= 1
        # Первая таблица — сводная: header + по строке на сценарий.
        assert len(doc.tables[0].rows) == 1 + len(scenarios)
    finally:
        if os.path.exists(path):
            os.unlink(path)


def test_build_report_empty_raises():
    with pytest.raises(ValueError):
        build_report([], tempfile.mktemp(suffix=".docx"))


def test_pick_best_prefers_higher_profit(scenarios):
    # Бизнес-класс выгоднее комфорта → должен быть рекомендован.
    idx, basis = _pick_best(scenarios)
    assert scenarios[idx][0] == "18 эт бизнес"
    assert "выгодност" in basis.lower()


def test_pick_best_returns_valid_index(scenarios):
    idx, _ = _pick_best(scenarios)
    assert 0 <= idx < len(scenarios)


# v0.11: отдельный отчёт по варианту (альбомный, подробный).

def test_variant_report_creates_landscape_docx(spb):
    from urban_model.export import build_variant_report
    from urban_model.models import FloorCluster
    s = Site(area_m2=130_000)
    tep = solve_max_kit(
        s,
        CalculationOptions(
            floors=12, planning_doc=True, residential_class="business",
            floor_clusters=[
                FloorCluster(area_m2=50_000, floors=9, label="Зона А"),
                FloorCluster(area_m2=80_000, floors=16, label="Зона Б"),
            ],
        ),
        spb,
    )
    path = tempfile.mktemp(suffix=".docx")
    try:
        build_variant_report("Максимум прибыли", tep, path)
        assert os.path.getsize(path) > 10000  # с диаграммами
        from docx import Document
        doc = Document(path)
        sec = doc.sections[0]
        # альбомная A4
        assert sec.page_width.cm > sec.page_height.cm
        assert round(sec.page_width.cm, 1) == 29.7
        assert len(doc.tables) >= 4  # KPI, баланс, жильё, парковки, кластеры, эконом
    finally:
        if os.path.exists(path):
            os.unlink(path)


def test_variant_report_without_economy(spb):
    from urban_model.export import build_variant_report
    s = Site(area_m2=50_000)
    tep = solve_max_kit(s, CalculationOptions(floors=12, include_economy=False), spb)
    assert tep.economy is None
    path = tempfile.mktemp(suffix=".docx")
    try:
        build_variant_report("Без экономики", tep, path)
        from docx import Document
        assert len(Document(path).tables) >= 3
    finally:
        if os.path.exists(path):
            os.unlink(path)


def test_comparison_report_landscape(scenarios):
    from urban_model.export import build_comparison_report
    path = tempfile.mktemp(suffix=".docx")
    try:
        build_comparison_report(scenarios, path)
        from docx import Document
        doc = Document(path)
        assert doc.sections[0].page_width.cm > doc.sections[0].page_height.cm
    finally:
        if os.path.exists(path):
            os.unlink(path)
